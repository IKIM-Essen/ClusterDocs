from pathlib import Path
import subprocess, shutil, re
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT=Path('/mnt/data/rcc-curriculum-v2')
SRC=ROOT/'source'
DOCX=ROOT/'docx'; PDF=ROOT/'pdf'
DOCX.mkdir(exist_ok=True); PDF.mkdir(exist_ok=True)
REFERENCE=Path('/mnt/data/rcc-onboarding/reference.docx')
BLUE='184E77'; LIGHT='EAF2F8'; DARK=RGBColor(0x20,0x2B,0x33); MUTED=RGBColor(0x5B,0x67,0x70); WHITE=RGBColor(255,255,255)

def shade(cell, fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=tcPr.find(qn('w:shd'))
    if shd is None: shd=OxmlElement('w:shd'); tcPr.append(shd)
    shd.set(qn('w:fill'),fill)

def margins(cell, top=70,start=75,bottom=70,end=75):
    tcPr=cell._tc.get_or_add_tcPr(); tcMar=tcPr.first_child_found_in('w:tcMar')
    if tcMar is None: tcMar=OxmlElement('w:tcMar'); tcPr.append(tcMar)
    for n,v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        x=tcMar.find(qn('w:'+n))
        if x is None: x=OxmlElement('w:'+n); tcMar.append(x)
        x.set(qn('w:w'),str(v)); x.set(qn('w:type'),'dxa')

def tr_flag(row, tag):
    trPr=row._tr.get_or_add_trPr(); el=OxmlElement(tag); trPr.append(el)

def set_repeat(row):
    trPr=row._tr.get_or_add_trPr(); el=OxmlElement('w:tblHeader'); el.set(qn('w:val'),'true'); trPr.append(el)

def set_cell_width(cell, width_cm):
    cell.width=Cm(width_cm)
    tcPr=cell._tc.get_or_add_tcPr(); tcW=tcPr.find(qn('w:tcW'))
    if tcW is None: tcW=OxmlElement('w:tcW'); tcPr.append(tcW)
    tcW.set(qn('w:w'),str(int(width_cm*567))); tcW.set(qn('w:type'),'dxa')

def postprocess(path: Path, part: int):
    doc=Document(path)
    sec=doc.sections[0]
    # Main content starts on a fresh page after title metadata.
    first_h1=True
    for p in doc.paragraphs:
        if p.style and p.style.name=='Heading 1' and first_h1:
            p.paragraph_format.page_break_before=True; first_h1=False
        if p.style and p.style.name.startswith('Heading'):
            p.paragraph_format.keep_with_next=True
    # Tables sized to text width.
    for table in doc.tables:
        table.style='Table Grid'; table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.autofit=False
        n=len(table.columns)
        if n==2: widths=[5.0,11.0]
        elif n==3: widths=[3.4,5.2,7.4]
        elif n==4: widths=[3.0,4.0,4.4,4.6]
        elif n==5: widths=[2.3,3.2,3.4,3.4,3.7]
        else: widths=[16.0/n]*n
        for ri,row in enumerate(table.rows):
            tr_flag(row,'w:cantSplit')
            if ri==0: set_repeat(row)
            for ci,cell in enumerate(row.cells):
                set_cell_width(cell,widths[min(ci,len(widths)-1)])
                cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
                margins(cell)
                shade(cell, BLUE if ri==0 else ('F5F8FA' if ri%2==0 else 'FFFFFF'))
                for p in cell.paragraphs:
                    # Pandoc uses an undefined Compact style in table cells; LibreOffice
                    # may render that text outside the table. Remove the style reference.
                    pPr=p._p.get_or_add_pPr()
                    pStyle=pPr.find(qn('w:pStyle'))
                    if pStyle is not None: pPr.remove(pStyle)
                    p.paragraph_format.space_after=Pt(0); p.paragraph_format.space_before=Pt(0); p.paragraph_format.line_spacing=1.0
                    for r in p.runs:
                        r.font.name='Liberation Sans'; r.font.size=Pt(7.9 if n>=4 else 8.4)
                        r.font.color.rgb=WHITE if ri==0 else DARK
                        if ri==0:r.font.bold=True
    # Header and footer.
    header=sec.header.paragraphs[0]
    header.text=f'RCC Onboarding - Part {part}'
    header.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    for r in header.runs:
        r.font.name='Liberation Sans'; r.font.size=Pt(8); r.font.color.rgb=MUTED
    footer=sec.footer.paragraphs[0]
    footer.clear(); footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=footer.add_run('IKIM RCC | Draft for technical and user review | ')
    r.font.name='Liberation Sans'; r.font.size=Pt(8); r.font.color.rgb=MUTED
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); footer._p.append(fld)
    doc.core_properties.title=f'RCC Onboarding - Part {part}'
    doc.core_properties.subject='Training material for biomedical researchers using the RCC cluster'
    doc.core_properties.author='IKIM RCC documentation proposal'
    doc.save(path)

for part in range(1,5):
    src=SRC/f'part{part}.md'
    out=DOCX/f'RCC_Onboarding_Part_{part}.docx'
    subprocess.run(['pandoc',str(src),'-f','gfm+yaml_metadata_block','-t','docx','--standalone',f'--reference-doc={REFERENCE}','-o',str(out)],check=True)
    postprocess(out,part)
    # Convert through the standard PDF helper.
    subprocess.run(['python','/home/oai/skills/pdfs/scripts/lo_convert_to_pdf.py',str(out),'--out_dir',str(PDF)],check=True)
print('Built DOCX and PDF files')
